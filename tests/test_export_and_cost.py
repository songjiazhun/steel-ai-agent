"""CAD / WPS 导出与费用核算关键测试。"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from typing import Any, Dict, List

import ezdxf
from docx import Document

from modules.cad_generator import CADGenerator
from modules.nesting_engine import NestingEngine2D, SteelPlateItem
from modules.pricing import build_quote_data, calc_total_weight_ton
from modules.wps_exporter import WPSExporter


class TestCostCalculation(unittest.TestCase):
    """材料重量与报价核算。"""

    def test_weight_formula(self) -> None:
        items: List[SteelPlateItem] = [
            SteelPlateItem("A", 1000.0, 1000.0, 10.0, 2, "Q235B"),
        ]
        # V = 1×1×0.01×2 = 0.02 m³ → 0.02×7.85 = 0.157 t
        weight: float = calc_total_weight_ton(items)
        self.assertAlmostEqual(weight, 0.02 * 7.85, places=6)

    def test_quote_includes_profit(self) -> None:
        items: List[SteelPlateItem] = [
            SteelPlateItem("A", 1000.0, 2000.0, 12.0, 1, "Q235B"),
        ]
        nesting: Dict[str, Any] = {
            "num_sheets": 2,
            "utilization_rate": 30.0,
            "stock_width_mm": 2200.0,
            "stock_length_mm": 6000.0,
            "sheets": [],
        }
        quote: Dict[str, Any] = build_quote_data(nesting, items)
        material: float = calc_total_weight_ton(items) * 4200.0
        process: float = 2 * 350.0
        expected: float = round((material + process) * (1.0 + 0.12), 2)
        self.assertEqual(quote["num_sheets"], 2)
        self.assertEqual(quote["process_cost"], process)
        self.assertEqual(quote["total_with_profit"], expected)
        self.assertGreaterEqual(len(quote["table_rows"]), 3)


class TestCADGenerator(unittest.TestCase):
    """DXF 图层与文件生成。"""

    def test_generate_dxf_layers_and_entities(self) -> None:
        engine = NestingEngine2D()
        items: List[SteelPlateItem] = [
            SteelPlateItem("A", 800.0, 1200.0, 12.0, 2, "Q235B"),
        ]
        nesting = engine.pack(items)

        with tempfile.TemporaryDirectory() as tmp:
            path: Path = Path(tmp) / "test.dxf"
            CADGenerator.generate_nesting_dxf(nesting, str(path))
            self.assertTrue(path.exists())
            self.assertGreater(path.stat().st_size, 0)

            doc = ezdxf.readfile(str(path))
            layer_names = {layer.dxf.name for layer in doc.layers}
            self.assertIn("0_STOCK_BORDER", layer_names)
            self.assertIn("0_PARTS", layer_names)
            self.assertIn("0_TEXT", layer_names)

            self.assertEqual(doc.layers.get("0_STOCK_BORDER").color, 1)
            self.assertEqual(doc.layers.get("0_PARTS").color, 3)
            self.assertEqual(doc.layers.get("0_TEXT").color, 2)

            msp = doc.modelspace()
            self.assertGreater(len(list(msp)), 0)


class TestWPSExporter(unittest.TestCase):
    """报价单 DOCX 结构。"""

    def test_generate_docx_title_and_table(self) -> None:
        quote_data: Dict[str, Any] = {
            "stock_width_mm": 2200.0,
            "stock_length_mm": 6000.0,
            "num_sheets": 1,
            "utilization_rate": 55.5,
            "material_cost": 1000.0,
            "process_cost": 350.0,
            "profit_rate": 0.12,
            "total_with_profit": 1512.0,
            "table_rows": [
                ["钢板 A", "800×1200×12 Q235B", "2 件", "4200/吨", "500.00"],
                ["加工费", "切割 × 1 张", "1 张", "350/张", "350.00"],
            ],
        }
        with tempfile.TemporaryDirectory() as tmp:
            path: Path = Path(tmp) / "quote.docx"
            WPSExporter.generate_wps_report(quote_data, str(path))
            self.assertTrue(path.exists())

            doc = Document(str(path))
            full_text: str = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("钢材套裁排版与智能化报价单", full_text)
            self.assertIn("一、排版分析", full_text)
            self.assertIn("二、结算明细", full_text)
            self.assertIn("55.50%", full_text)
            self.assertIn("1,512.00", full_text)

            self.assertEqual(len(doc.tables), 1)
            table = doc.tables[0]
            self.assertEqual(len(table.columns), 5)
            self.assertEqual(table.rows[0].cells[0].text, "项目")
            self.assertEqual(len(table.rows), 3)  # 表头 + 2 数据行


if __name__ == "__main__":
    unittest.main()
