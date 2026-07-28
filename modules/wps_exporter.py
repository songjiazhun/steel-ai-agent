"""WPS 商务报价单导出模块（python-docx，兼容 WPS Office）。"""

from __future__ import annotations

from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class WPSExporter:
    """生成可在 WPS Office 完美打开的钢材套裁报价单。"""

    HEADER_BG: str = "0F4C81"
    TITLE_COLOR: RGBColor = RGBColor(0x0F, 0x4C, 0x81)
    TOTAL_COLOR: RGBColor = RGBColor(0xCC, 0x00, 0x00)

    @staticmethod
    def generate_wps_report(quote_data: dict, output_path: str) -> None:
        """根据报价数据生成 WPS 兼容的 DOCX 报价单。

        Args:
            quote_data: 报价数据字典，需包含排版与费用明细字段。
            output_path: DOCX 文件输出路径。
        """
        doc: Document = Document()
        WPSExporter._set_page_margins(doc, inches=0.8)
        WPSExporter._add_title(doc)
        WPSExporter._add_nesting_analysis(doc, quote_data)
        WPSExporter._add_settlement_table(doc, quote_data)
        WPSExporter._add_total(doc, quote_data)
        doc.save(output_path)

    @staticmethod
    def _set_page_margins(doc: Document, inches: float) -> None:
        """设置页边距（英寸）。"""
        for section in doc.sections:
            section.top_margin = Inches(inches)
            section.bottom_margin = Inches(inches)
            section.left_margin = Inches(inches)
            section.right_margin = Inches(inches)

    @staticmethod
    def _add_title(doc: Document) -> None:
        """添加居中主标题。"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("钢材套裁排版与智能化报价单")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = WPSExporter.TITLE_COLOR
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    @staticmethod
    def _add_nesting_analysis(doc: Document, quote_data: dict) -> None:
        """添加一、排版分析段落。"""
        heading = doc.add_paragraph()
        h_run = heading.add_run("一、排版分析")
        h_run.bold = True
        h_run.font.size = Pt(14)

        stock_w: float = float(quote_data.get("stock_width_mm", 2200.0))
        stock_l: float = float(quote_data.get("stock_length_mm", 6000.0))
        num_sheets: int = int(quote_data.get("num_sheets", 0))
        utilization: float = float(quote_data.get("utilization_rate", 0.0))

        lines: List[str] = [
            f"母板规格：{stock_w:.0f} mm × {stock_l:.0f} mm",
            f"用板张数：{num_sheets} 张",
        ]
        for line in lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(4)

        util_p = doc.add_paragraph()
        util_p.paragraph_format.space_after = Pt(8)
        util_label = util_p.add_run("综合套裁利用率：")
        util_label.font.size = Pt(11)
        util_value = util_p.add_run(f"{utilization:.2f}%")
        util_value.bold = True
        util_value.font.size = Pt(12)
        util_value.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    @staticmethod
    def _add_settlement_table(doc: Document, quote_data: dict) -> None:
        """添加二、结算明细表格（5 列，表头 XML 背景色）。"""
        heading = doc.add_paragraph()
        h_run = heading.add_run("二、结算明细")
        h_run.bold = True
        h_run.font.size = Pt(14)

        headers: List[str] = ["项目", "规格/说明", "数量", "单价(元)", "金额(元)"]
        rows: List[List[str]] = quote_data.get("table_rows", [])

        table = doc.add_table(rows=1 + len(rows), cols=5)
        table.style = "Table Grid"

        # 表头
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            WPSExporter._set_cell_shading(cell, WPSExporter.HEADER_BG)

        # 数据行
        for row_idx, row_data in enumerate(rows):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if col_idx >= 2
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                run = p.add_run(str(value))
                run.font.size = Pt(10)

    @staticmethod
    def _add_total(doc: Document, quote_data: dict) -> None:
        """结尾右对齐输出最终含税与利润的总报价。"""
        doc.add_paragraph()
        total_amount: float = float(quote_data.get("total_with_profit", 0.0))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"最终报价（含目标利润）：¥ {total_amount:,.2f}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = WPSExporter.TOTAL_COLOR

        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        n_run = note.add_run(
            f"（材料费 ¥{float(quote_data.get('material_cost', 0)): ,.2f}"
            f" + 加工费 ¥{float(quote_data.get('process_cost', 0)): ,.2f}"
            f" + 利润 {float(quote_data.get('profit_rate', 0.12)) * 100:.0f}%）"
        )
        n_run.font.size = Pt(9)
        n_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    @staticmethod
    def _set_cell_shading(cell: Any, hex_color: str) -> None:
        """通过 XML 设置单元格背景色，保证 WPS 兼容。"""
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)
